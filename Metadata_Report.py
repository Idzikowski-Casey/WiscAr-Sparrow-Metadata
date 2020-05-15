from Statisitcal_Analysis import *

from docx import Document

doc = Document()
doc.add_heading('MetaData Report', 0)

r = str(len(Uploadable))
first = doc.add_paragraph()
first.add_run(str(Percent_Loc_Found_Metada)).bold = True
first.add_run('%').bold = True
first.add_run(' of the literature search samples that contain at least location metadata.')

paragraph2 = doc.add_paragraph()
paragraph2.add_run(str(Percent_Loc_Found_Online)).bold = True
paragraph2.add_run('%').bold = True
paragraph2.add_run(' of Sparrow samples that contain at least location metadata.')

second = doc.add_paragraph()
second.add_run('There are ')
second.add_run(r).bold = True
second.add_run(' samples with updated metadata that need to be pushed to sparrow. These are samples that '
               'metadata has been found since the original literature search. There is an updated JSON object for '
               'pushing to sparrow. The JSON file can be found in ')
second.add_run('upload.json').bold = True
second.add_run('.')

paratwoB = doc.add_paragraph()
paratwoB.add_run('There are ')
paratwoB.add_run(str(len(MetaMissing))).bold = True
paratwoB.add_run(' samples with missing metadata still. They have been exported to an excel sheet. '
                 'They can be found in')
paratwoB.add_run('MissingMetadata.xlsx').bold = True
paratwoB.add_run('.')

third = doc.add_paragraph('There is a simplified spread sheet for all samples called ')
third.add_run('Camparison_Sheet.xls').bold = True
third.add_run('.')

fourth = doc.add_paragraph()
fourth.add_run('online_metadata.xls').bold = True
fourth.add_run(' contains all the samples and their metadata from sparrow.')


doc.save('Metadata_Report.docx')



